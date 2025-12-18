import os
import time
import sys
import shutil
import json
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QLineEdit, QPushButton, QFileDialog, QTableWidget,
                             QTableWidgetItem, QMessageBox, QGroupBox, QGridLayout,
                             QHeaderView, QTextEdit, QCheckBox, QComboBox)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont
from openpyxl import load_workbook
from docx import Document
from docxcompose.composer import Composer

class WorkerThread(QThread):
    """工作线程，用于处理耗时操作，避免UI卡顿"""
    status_updated = pyqtSignal(str)
    finished = pyqtSignal(int)

    def __init__(self, generator, operation):
        super().__init__()
        self.generator = generator
        self.operation = operation  # 'generate' 或 'merge'

    def run(self):
        try:
            if self.operation == 'generate':
                # 生成证书前先清理输出目录中的Word文件
                self.status_updated.emit("清理输出目录中的原有Word文件...")
                if os.path.exists(self.generator.output_dir):
                    for filename in os.listdir(self.generator.output_dir):
                        if filename.endswith(".docx") and not filename.startswith("~$"):
                            file_path = os.path.join(self.generator.output_dir, filename)
                            try:
                                os.remove(file_path)
                                self.status_updated.emit(f"已删除: {filename}")
                            except Exception as e:
                                self.status_updated.emit(f"删除{filename}失败: {str(e)}")
                
                # 生成证书
                result = self.generator.generate_documents(self.status_updated.emit)
                
                # 自动合并证书
                if result == 0:
                    self.status_updated.emit("\n开始自动合并证书...")
                    merge_result = self.generator.merge_docx(self.status_updated.emit)
                    self.finished.emit(0 if merge_result else 1)
                else:
                    self.finished.emit(1)
            elif self.operation == 'merge':
                result = self.generator.merge_docx(self.status_updated.emit)
                self.finished.emit(0 if result else 1)
        except Exception as e:
            self.status_updated.emit(f"操作出错: {str(e)}")
            self.finished.emit(1)

class DocumentGenerator:
    def __init__(self, database, template, doc_output, merge_output, replacement_config):
        self.excel_file = database
        self.template_file = template
        self.output_dir = doc_output
        self.merge_output = merge_output
        self.replacement_config = replacement_config
        self.generated_files = [] # 已生成的Word文件列表
        
        # 确保输出目录存在
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def read_excel_data(self, status_callback):
        try:
            excel_start = time.perf_counter()
            workbook = load_workbook(self.excel_file, data_only=True)
            sheet = workbook.active
            status_callback(f"成功加载数据库文件：{self.excel_file}（工作表名：{sheet.title}）")

            headers = [cell.value.strip() if cell.value else "" for cell in sheet[1]]
            headers = [h for h in headers if h]
            
            required_headers = [item["excel_header"] for item in self.replacement_config]
            missing_headers = [h for h in required_headers if h not in headers]
            if missing_headers:
                raise ValueError(f"Excel缺少必要表头：{missing_headers}")

            data = []
            for row_num, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
                row_data = {}
                for config in self.replacement_config:
                    header = config["excel_header"]
                    col_index = headers.index(header)
                    cell_value = row[col_index] if col_index < len(row) else ""
                    row_data[header] = str(cell_value).strip() if cell_value else ""
                
                # 必填字段校验
                required_empty = [cfg["excel_header"] for cfg in self.replacement_config 
                                 if cfg.get("required", False) and not row_data[cfg["excel_header"]]]
                if required_empty:
                    status_callback(f"跳过不完整数据行（第{row_num}行）：必填字段{required_empty}为空")
                else:
                    data.append(row_data)

            workbook.close()
            excel_end = time.perf_counter()
            status_callback(f"成功读取 {len(data)} 条数据（耗时：{excel_end - excel_start:.2f}秒）")
            return data

        except FileNotFoundError:
            status_callback(f"错误：未找到Excel文件「{self.excel_file}」")
            return []
        except Exception as e:
            status_callback(f"读取Excel错误：{str(e)}")
            return []

    def replace_placeholder_in_run(self, run_text, row_data):
        if not run_text:
            return run_text
        
        replaced_text = run_text
        for config in self.replacement_config:
            placeholder = config["placeholder"]
            excel_header = config["excel_header"]
            cell_value = row_data.get(excel_header, "")
            
            # 核心逻辑：空值删除占位符，非空按格式替换
            if not cell_value:  # 数据库值为空 → 删除占位符文本
                replaced_text = replaced_text.replace(placeholder, "")
            else:  # 数据库值非空 → 按format格式化替换
                format_str = config.get("format", "{0}")  # 默认直接替换
                replaced_text = replaced_text.replace(placeholder, format_str.format(cell_value))
        return replaced_text

    def process_paragraph(self, paragraph, row_data):
        if not paragraph.text.strip():
            return
        
        for run in paragraph.runs:
            original_run_text = run.text
            if not original_run_text:
                continue
            
            # 检查是否包含任意占位符
            placeholders = [item["placeholder"] for item in self.replacement_config]
            has_placeholder = any(ph in original_run_text for ph in placeholders)
            if not has_placeholder:
                continue
            
            # 替换/删除占位符（保留格式）
            new_run_text = self.replace_placeholder_in_run(original_run_text, row_data)
            if new_run_text != original_run_text:
                run.text = new_run_text

    def replace_placeholders(self, doc, row_data):
        # 处理普通段落
        for para in doc.paragraphs:
            self.process_paragraph(para, row_data)
        
        # 处理表格
        if doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.process_paragraph(para, row_data)

    def clean_filename(self, filename):
        """清理文件名中的非法字符"""
        invalid_chars = '/\\:*?"<>|' # Windows系统非法字符
        for char in invalid_chars:
            filename = filename.replace(char, '_') # 用下划线替换非法字符
        # 额外处理首尾空格/换行，避免文件名异常
        filename = filename.strip()
        return filename

    def generate_documents(self, status_callback):
        total_start = time.perf_counter()
        data = self.read_excel_data(status_callback)
        
        if not data:
            status_callback("\n❌ 无有效数据，无法生成证书")
            return 1
        
        # 清空之前的记录
        self.generated_files = []

        total = len(data)
        success_count = 0
        fail_count = 0
        
        # 获取文件名字段
        filename_fields = [cfg["excel_header"] for cfg in self.replacement_config if cfg.get("use_in_filename", False)]
        if not filename_fields:
            filename_fields = [self.replacement_config[0]["excel_header"]]
        
        for index, item in enumerate(data, start=1):
            try:
                doc = Document(self.template_file)  
                self.replace_placeholders(doc, item)
                
                # 生成文件名
                filename_parts = [str(item[field]) for field in filename_fields if item[field]]
                raw_filename = "_".join(filename_parts) + ".docx"
                filename = self.clean_filename(raw_filename)  # 清理非法字符
                file_path = os.path.join(self.output_dir, filename)
                doc.save(file_path)
                
                # 记录生成的文件路径（按数据库顺序）
                self.generated_files.append(file_path)

                success_count += 1
                status_callback(f"已生成：{filename}（{index}/{total}）")
            except Exception as e:
                fail_count += 1
                if "No such file or directory" in str(e) and "docx" in str(e):
                    status_callback(f"处理{item.get(filename_fields[0], '')}时出错：文件名含非法字符，{str(e)}")
                else:
                    status_callback(f"处理{item.get(filename_fields[0], '')}时出错：{str(e)}")
        
        # 输出汇总
        total_end = time.perf_counter()
        total_elapsed = total_end - total_start
        status_callback(f"生成完成：{success_count}/{success_count + fail_count}")
        status_callback(f"总耗时长：{total_elapsed:.2f}秒（平均：{total_elapsed/total:.2f}个/秒）" if total else f"总耗时长：{total_elapsed:.2f}秒")
        status_callback(f"保存路径：{os.path.abspath(self.output_dir)}")

        return 0 if success_count > 0 else 1

    def merge_docx(self, status_callback):
        docx_paths = self.generated_files.copy()
        
        # 如果没有记录，则回退到原有方式
        if not docx_paths:
            status_callback("警告：未找到生成记录，将按文件名排序合并（可能与数据库顺序不一致）")
            for filename in os.listdir(self.output_dir):
                if filename.endswith(".docx") and not filename.startswith("~$"):
                    docx_paths.append(os.path.abspath(os.path.join(self.output_dir, filename)))
            docx_paths.sort()
        
        if not docx_paths:
            status_callback("错误：未找到有效docx文件！")
            return None

        main_doc = Document(docx_paths[0])
        composer = Composer(main_doc)
        main_section = main_doc.sections[0]
        main_margins = (main_section.left_margin, main_section.right_margin, main_section.top_margin, main_section.bottom_margin)
        main_page_size = (main_section.page_width, main_section.page_height)

        total_docs = len(docx_paths)
        status_callback(f"开始合并 {total_docs} 个文档...")
        
        for i, doc_path in enumerate(docx_paths[1:], 1):
            try:
                sub_doc = Document(doc_path)
                for section in sub_doc.sections:
                    section.left_margin, section.right_margin = main_margins[0], main_margins[1]
                    section.top_margin, section.bottom_margin = main_margins[2], main_margins[3]
                    section.page_width, section.page_height = main_page_size[0], main_page_size[1]

                main_doc.add_page_break()
                composer.append(sub_doc)
                status_callback(f"已合并 {i+1}/{total_docs} 个文档")
            except Exception as e:
                status_callback(f"合并 {os.path.basename(doc_path)} 时出错：{str(e)}")

        composer.save(self.merge_output)
        status_callback(f"\n合并完成！")
        status_callback(f"保存路径：{os.path.abspath(self.merge_output)}")
        return self.merge_output

class CertificateGeneratorGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        # 预设配置数据
        self.presets = {
            "就业创业培训": [
                {"excel_header": "姓名", "placeholder": "湖小招", "format": "{0}", "required": True, "use_in_filename": True},
                {"excel_header": "学号", "placeholder": "14242300000", "format": "{0}", "required": True, "use_in_filename": True},
                {"excel_header": "项目", "placeholder": "XXXX", "format": "{0}", "required": True, "use_in_filename": False},
                {"excel_header": "学时", "placeholder": "（学时）", "format": "（{0}学时）", "required": True, "use_in_filename": False},
            ],
            "职业规划大赛": [
                {"excel_header": "姓名", "placeholder": "湖小招", "format": "{0}", "required": True, "use_in_filename": True},
                {"excel_header": "赛道", "placeholder": "XXXX", "format": "{0}", "required": True, "use_in_filename": False},
                {"excel_header": "奖项", "placeholder": "特等奖", "format": "{0}", "required": True, "use_in_filename":False},
                {"excel_header": "指导老师", "placeholder": "指导老师：", "format": "指导老师：{0}", "required": False, "use_in_filename": False},
                {"excel_header": "团队成员", "placeholder": "团队成员：", "format": "团队成员：{0}", "required": False, "use_in_filename": False}
            ],
            "金种子": [
                {"excel_header": "项目名称", "placeholder": "AAA", "format": "{0}", "required": True, "use_in_filename": True},
                {"excel_header": "赛道", "placeholder": "XXXX", "format": "{0}", "required": True, "use_in_filename": False},
                {"excel_header": "奖项", "placeholder": "特等奖", "format": "{0}", "required": True, "use_in_filename": False},
                {"excel_header": "团队成员", "placeholder": "团队成员：", "format": "团队成员：{0}", "required": True, "use_in_filename": False},
                {"excel_header": "指导老师", "placeholder": "指导老师：", "format": "指导老师：{0}", "required": False, "use_in_filename": False}
            ],
            "自定义": []  # 初始为空，用户自定义时保存
        }
        self.initUI()
        self.worker = None
        
    def initUI(self):
        # 设置窗口基本属性
        self.setWindowTitle('证书生成器')
        self.setGeometry(100, 100, 1000, 1000)
        
        # 创建中心部件和主布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 创建路径配置区域
        path_group = QGroupBox("文件路径配置")
        path_layout = QGridLayout()
        
        # Excel数据库文件
        path_layout.addWidget(QLabel("数据库（Excel文件）:"), 0, 0)
        self.excel_path = QLineEdit("./数据库.xlsx")
        path_layout.addWidget(self.excel_path, 0, 1)
        self.btn_excel = QPushButton("浏览...")
        self.btn_excel.clicked.connect(self.select_excel)
        path_layout.addWidget(self.btn_excel, 0, 2)
        
        # Word模板文件
        path_layout.addWidget(QLabel("模板（Word文件）:"), 1, 0)
        self.template_path = QLineEdit("./模板.docx")
        path_layout.addWidget(self.template_path, 1, 1)
        self.btn_template = QPushButton("浏览...")
        self.btn_template.clicked.connect(self.select_template)
        path_layout.addWidget(self.btn_template, 1, 2)
        
        # 输出目录
        path_layout.addWidget(QLabel("输出目录:"), 2, 0)
        self.output_dir = QLineEdit("./生成的证书（未合并）")
        path_layout.addWidget(self.output_dir, 2, 1)
        self.btn_output = QPushButton("浏览...")
        self.btn_output.clicked.connect(self.select_output_dir)
        path_layout.addWidget(self.btn_output, 2, 2)
        
        # 合并文件
        path_layout.addWidget(QLabel("合并文件:"), 3, 0)
        self.merge_file = QLineEdit("./生成的证书.docx")
        path_layout.addWidget(self.merge_file, 3, 1)
        self.btn_merge = QPushButton("浏览...")
        self.btn_merge.clicked.connect(self.select_merge_file)
        path_layout.addWidget(self.btn_merge, 3, 2)
        
        path_group.setLayout(path_layout)
        main_layout.addWidget(path_group)
        
        # 创建替换配置区域
        replace_group = QGroupBox("替换配置")
        replace_layout = QVBoxLayout()
        
        # 预设选择区域
        preset_layout = QHBoxLayout()
        preset_layout.addWidget(QLabel("选择预设:"))
        self.preset_combobox = QComboBox()
        self.preset_combobox.addItems(list(self.presets.keys()))
        self.preset_combobox.currentIndexChanged.connect(self.on_preset_changed)
        preset_layout.addWidget(self.preset_combobox)
        preset_layout.addStretch()
        replace_layout.addLayout(preset_layout)
        
        # 表格控件
        self.replace_table = QTableWidget()
        self.replace_table.setColumnCount(5)
        self.replace_table.setHorizontalHeaderLabels(["数据库类别", "替换的模板字符", "替换后的格式", "必填", "以此命名文件"])
        # 设置列宽自适应
        header = self.replace_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.Stretch)
        header.setSectionResizeMode(2, QHeaderView.Stretch)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.ResizeToContents)
        
        replace_layout.addWidget(self.replace_table)
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        self.btn_add = QPushButton("添加")
        self.btn_add.clicked.connect(self.add_row)
        self.btn_remove = QPushButton("删除选中")
        self.btn_remove.clicked.connect(self.remove_row)
        self.btn_import = QPushButton("导入配置")
        self.btn_import.clicked.connect(self.import_config)
        self.btn_export = QPushButton("导出配置")
        self.btn_export.clicked.connect(self.export_config)
        
        btn_layout.addWidget(self.btn_add)
        btn_layout.addWidget(self.btn_remove)
        btn_layout.addWidget(self.btn_import)
        btn_layout.addWidget(self.btn_export)
        
        replace_layout.addLayout(btn_layout)
        replace_group.setLayout(replace_layout)
        main_layout.addWidget(replace_group)
        
        # 初始加载第一个预设
        self.load_preset_config(list(self.presets.keys())[0])
        
        # 创建操作按钮区域
        btn_group = QHBoxLayout()
        self.btn_generate = QPushButton("生成并合并证书")
        self.btn_generate.setMinimumHeight(40)
        self.btn_generate.clicked.connect(self.generate_and_merge_certificates)
        
        self.btn_clear_status = QPushButton("清空状态信息")
        self.btn_clear_status.setMinimumHeight(40)
        self.btn_clear_status.clicked.connect(self.clear_status)
        
        btn_group.addWidget(self.btn_generate)
        btn_group.addWidget(self.btn_clear_status)
        main_layout.addLayout(btn_group)
        
        # 创建状态输出区域
        status_group = QGroupBox("状态信息")
        status_layout = QVBoxLayout()
        self.status_text = QTextEdit()
        self.status_text.setReadOnly(True)
        status_layout.addWidget(self.status_text)
        status_group.setLayout(status_layout)
        main_layout.addWidget(status_group)
        
        # 设置拉伸因子，让状态区域可以伸缩
        main_layout.setStretch(0, 1)  # 路径配置
        main_layout.setStretch(1, 3)  # 替换配置
        main_layout.setStretch(2, 1)  # 按钮区域
        main_layout.setStretch(3, 2)  # 状态信息
        
        self.show()
    
    def load_preset_config(self, preset_name):
        """加载指定的预设配置到表格中"""
        # 先清空表格
        while self.replace_table.rowCount() > 0:
            self.replace_table.removeRow(0)
        
        # 加载预设配置
        configs = self.presets.get(preset_name, [])
        for config in configs:
            self.add_row(config)
    
    def on_preset_changed(self, index):
        """当预设选择改变时触发"""
        preset_name = self.preset_combobox.currentText()
        # 如果当前是自定义配置且有内容，先保存
        if self.preset_combobox.itemText(index-1) == "自定义" and self.replace_table.rowCount() > 0:
            self.save_custom_preset()
        
        self.load_preset_config(preset_name)
    
    def save_custom_preset(self):
        """保存当前配置为自定义预设"""
        custom_config = []
        for row in range(self.replace_table.rowCount()):
            excel_header_item = self.replace_table.item(row, 0)
            placeholder_item = self.replace_table.item(row, 1)
            
            if excel_header_item and placeholder_item and excel_header_item.text().strip() and placeholder_item.text().strip():
                custom_config.append({
                    "excel_header": excel_header_item.text().strip(),
                    "placeholder": placeholder_item.text().strip(),
                    "format": self.replace_table.item(row, 2).text().strip() if self.replace_table.item(row, 2) else "{0}",
                    "required": self.replace_table.item(row, 3).checkState() == Qt.Checked,
                    "use_in_filename": self.replace_table.item(row, 4).checkState() == Qt.Checked
                })
        
        self.presets["自定义"] = custom_config
    
    def add_row(self, config=None):
        row_pos = self.replace_table.rowCount()
        self.replace_table.insertRow(row_pos)
        
        # 如果有配置数据，则使用配置数据，否则使用默认值
        if config:
            excel_header = config.get("excel_header", "")
            placeholder = config.get("placeholder", "")
            format_str = config.get("format", "{0}")
            required = config.get("required", False)
            use_in_filename = config.get("use_in_filename", False)
        else:
            excel_header = ""
            placeholder = ""
            format_str = "{0}"
            required = False
            use_in_filename = False
        
        # 创建单元格并设置值
        self.replace_table.setItem(row_pos, 0, QTableWidgetItem(excel_header))
        self.replace_table.setItem(row_pos, 1, QTableWidgetItem(placeholder))
        self.replace_table.setItem(row_pos, 2, QTableWidgetItem(format_str))
        
        # 创建复选框单元格
        required_checkbox = QTableWidgetItem()
        required_checkbox.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled)
        required_checkbox.setCheckState(Qt.Checked if required else Qt.Unchecked)
        self.replace_table.setItem(row_pos, 3, required_checkbox)
        
        filename_checkbox = QTableWidgetItem()
        filename_checkbox.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled)
        filename_checkbox.setCheckState(Qt.Checked if use_in_filename else Qt.Unchecked)
        self.replace_table.setItem(row_pos, 4, filename_checkbox)
    
    def remove_row(self):
        selected_rows = set()
        for item in self.replace_table.selectedItems():
            selected_rows.add(item.row())
        
        if not selected_rows:
            QMessageBox.warning(self, "警告", "请先选择要删除的行")
            return
        
        # 从最大的行号开始删除，避免索引问题
        for row in sorted(selected_rows, reverse=True):
            self.replace_table.removeRow(row)
        
        # 如果当前是自定义预设，更新保存
        if self.preset_combobox.currentText() == "自定义":
            self.save_custom_preset()
    
    def import_config(self):
        """导入配置文件"""
        filename, _ = QFileDialog.getOpenFileName(self, "导入配置", "", "配置文件 (*.json)")
        if not filename:
            return
        
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                configs = json.load(f)
            
            # 验证配置格式
            if not isinstance(configs, list):
                raise ValueError("配置文件格式不正确")
            
            # 清空当前表格并加载导入的配置
            while self.replace_table.rowCount() > 0:
                self.replace_table.removeRow(0)
            
            for config in configs:
                # 验证必要字段
                if not all(key in config for key in ["excel_header", "placeholder"]):
                    raise ValueError("配置文件缺少必要字段")
                
                self.add_row(config)
            
            # 切换到自定义预设并保存
            self.preset_combobox.setCurrentText("自定义")
            self.save_custom_preset()
            self.log_status(f"成功导入配置文件: {filename}")
            
        except Exception as e:
            QMessageBox.critical(self, "导入失败", f"导入配置时出错: {str(e)}")
    
    def export_config(self):
        """导出当前配置"""
        # 获取当前配置
        configs = self.get_replacement_config()
        if not configs:
            QMessageBox.warning(self, "警告", "没有可导出的有效配置")
            return
        
        filename, _ = QFileDialog.getSaveFileName(self, "导出配置", "", "配置文件 (*.json)")
        if not filename:
            return
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(configs, f, ensure_ascii=False, indent=2)
            
            self.log_status(f"成功导出配置文件: {filename}")
            
            # 如果当前是自定义预设，更新保存
            if self.preset_combobox.currentText() == "自定义":
                self.save_custom_preset()
                
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出配置时出错: {str(e)}")
    
    def get_replacement_config(self):
        config = []
        for row in range(self.replace_table.rowCount()):
            excel_header_item = self.replace_table.item(row, 0)
            placeholder_item = self.replace_table.item(row, 1)
            
            # 验证必要字段
            if not excel_header_item or not excel_header_item.text().strip():
                QMessageBox.warning(self, "警告", f"第{row+1}行的Excel表头不能为空")
                return None
            
            if not placeholder_item or not placeholder_item.text().strip():
                QMessageBox.warning(self, "警告", f"第{row+1}行的占位符不能为空")
                return None
            
            # 收集数据
            config.append({
                "excel_header": excel_header_item.text().strip(),
                "placeholder": placeholder_item.text().strip(),
                "format": self.replace_table.item(row, 2).text().strip() if self.replace_table.item(row, 2) else "{0}",
                "required": self.replace_table.item(row, 3).checkState() == Qt.Checked,
                "use_in_filename": self.replace_table.item(row, 4).checkState() == Qt.Checked
            })
        
        return config
    
    def select_excel(self):
        filename, _ = QFileDialog.getOpenFileName(self, "选择Excel数据库", "", "Excel文件 (*.xlsx *.xls)")
        if filename:
            self.excel_path.setText(filename)
    
    def select_template(self):
        filename, _ = QFileDialog.getOpenFileName(self, "选择模板文件", "", "Word文件 (*.docx)")
        if filename:
            self.template_path.setText(filename)
    
    def select_output_dir(self):
        dirname = QFileDialog.getExistingDirectory(self, "选择输出目录", "")
        if dirname:
            self.output_dir.setText(dirname)
    
    def select_merge_file(self):
        filename, _ = QFileDialog.getSaveFileName(self, "保存合并文件", "", "Word文件 (*.docx)")
        if filename:
            self.merge_file.setText(filename)
    
    def log_status(self, message):
        """实时输出状态信息"""
        self.status_text.append(message)
        # 滚动到底部
        self.status_text.moveCursor(self.status_text.textCursor().End)
        # 强制刷新UI
        QApplication.processEvents()
    
    def clear_status(self):
        """清空状态信息"""
        self.status_text.clear()
    
    def generate_and_merge_certificates(self):
        """生成并合并证书（一步完成）"""
        # 清空状态
        self.clear_status()
        
        # 如果当前是自定义预设，先保存
        if self.preset_combobox.currentText() == "自定义":
            self.save_custom_preset()
        
        # 获取配置
        replacement_config = self.get_replacement_config()
        if not replacement_config:
            return
        
        # 获取路径
        excel_file = self.excel_path.text().strip()
        template_file = self.template_path.text().strip()
        output_dir = self.output_dir.text().strip()
        merge_file = self.merge_file.text().strip()
        
        # 验证文件存在
        if not os.path.exists(excel_file):
            QMessageBox.critical(self, "错误", f"Excel文件不存在: {excel_file}")
            return
        
        if not os.path.exists(template_file):
            QMessageBox.critical(self, "错误", f"模板文件不存在: {template_file}")
            return
        
        # 禁用生成按钮，防止重复点击
        self.btn_generate.setEnabled(False)
        
        # 创建生成器和工作线程
        try:
            generator = DocumentGenerator(excel_file, template_file, output_dir, merge_file, replacement_config)
            self.worker = WorkerThread(generator, 'generate')
            self.worker.status_updated.connect(self.log_status)
            self.worker.finished.connect(self.on_operation_finished)
            self.log_status("开始生成并合并证书...")
            self.worker.start()
        except Exception as e:
            self.log_status(f"初始化出错: {str(e)}")
            self.btn_generate.setEnabled(True)
    
    def on_operation_finished(self, result):
        """操作完成后的回调"""
        self.btn_generate.setEnabled(True)
        if result == 0:
            QMessageBox.information(self, "成功", "证书生成和合并完成！")
        else:
            QMessageBox.warning(self, "警告", "操作过程中出现错误")
        self.worker = None

if __name__ == "__main__":
    app = QApplication(sys.argv)
    # 确保中文显示正常
    font = QFont("SimHei")
    app.setFont(font)
    window = CertificateGeneratorGUI()
    sys.exit(app.exec_())